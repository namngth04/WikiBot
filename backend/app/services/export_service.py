import os
import re
from datetime import datetime
from typing import List, Dict, Any, Tuple
import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# XML Helper Functions for MS Word Styling
# ==========================================

def add_p_border_bottom(paragraph):
    """Adds a thin bottom border to a paragraph (used for header line)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 1/8 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')  # Light gray
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_p_border_top(paragraph):
    """Adds a thin top border to a paragraph (used for footer line)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')  # 1/8 pt
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'CCCCCC')  # Light gray
    pBdr.append(top)
    pPr.append(pBdr)

def add_page_number(run):
    """Injects a dynamic Page field character inside a run"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# ==========================================
# Markdown Table Parsing Helpers
# ==========================================

def is_markdown_table_line(line: str) -> bool:
    """Checks if a single line is part of a markdown table"""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|')

def parse_markdown_table(lines: List[str]) -> Tuple[List[str], List[List[str]]]:
    """Parses list of raw markdown lines into (headers, rows)"""
    headers = []
    rows = []
    
    # Filter out separating lines like |---|---|
    table_lines = []
    for l in lines:
        stripped = l.strip()
        # Skip separating rows
        if re.match(r'^\|[\s\-\|:]+\|$', stripped):
            continue
        table_lines.append(stripped)
        
    if not table_lines:
        return headers, rows
        
    # Process header
    header_raw = table_lines[0]
    headers = [cell.strip() for cell in header_raw.split('|')[1:-1]]
    
    # Process rows
    for row_raw in table_lines[1:]:
        row_cells = [cell.strip() for cell in row_raw.split('|')[1:-1]]
        # Pad row cells if it has fewer elements than headers
        while len(row_cells) < len(headers):
            row_cells.append("")
        rows.append(row_cells[:len(headers)])
        
    return headers, rows

def extract_content_blocks(text: str) -> List[Dict[str, Any]]:
    """Splits text into blocks of either 'text' or 'table'"""
    blocks = []
    lines = text.split('\n')
    
    in_table = False
    current_table_lines = []
    current_text_lines = []
    
    for line in lines:
        if is_markdown_table_line(line):
            if not in_table:
                # Flush text lines
                if current_text_lines:
                    blocks.append({
                        "type": "text",
                        "content": "\n".join(current_text_lines)
                    })
                    current_text_lines = []
                in_table = True
            current_table_lines.append(line)
        else:
            if in_table:
                # Flush table lines
                if current_table_lines:
                    blocks.append({
                        "type": "table",
                        "content": current_table_lines
                    })
                    current_table_lines = []
                in_table = False
            current_text_lines.append(line)
            
    # Flush remaining
    if in_table and current_table_lines:
        blocks.append({
            "type": "table",
            "content": current_table_lines
        })
    elif current_text_lines:
        blocks.append({
            "type": "text",
            "content": "\n".join(current_text_lines)
        })
        
    return blocks

# ==========================================
# Core Export Logic
# ==========================================

class ExportService:
    @staticmethod
    def export_to_txt(conversation_title: str, messages: List[Dict[str, Any]]) -> str:
        """Exports conversation history to a clean TXT format"""
        output = []
        output.append(f"==================================================")
        output.append(f"WIKIBOT CONVERSATION REPORT")
        output.append(f"Tiêu đề: {conversation_title}")
        output.append(f"Thời gian xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        output.append(f"==================================================\n")
        
        for msg in messages:
            role_label = "NGƯỜI DÙNG" if msg.get("role") == "user" else "WIKIBOT (ASSISTANT)"
            output.append(f"[{role_label}]:")
            output.append(msg.get("content", "").strip())
            output.append("\n" + "-"*40 + "\n")
            
        return "\n".join(output)

    @staticmethod
    def export_to_markdown(conversation_title: str, messages: List[Dict[str, Any]]) -> str:
        """Exports conversation history to GitHub Flavored Markdown"""
        output = []
        output.append(f"# Báo cáo cuộc hội thoại: {conversation_title}")
        output.append(f"*Thời gian xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}*\n")
        output.append("---")
        
        for msg in messages:
            if msg.get("role") == "user":
                output.append(f"### 👤 Người dùng")
                output.append(f"> {msg.get('content', '').strip()}\n")
            else:
                output.append(f"### 🤖 WikiBot")
                output.append(f"{msg.get('content', '').strip()}\n")
                output.append("---")
                
        return "\n".join(output)

    @staticmethod
    def export_to_docx(conversation_title: str, messages: List[Dict[str, Any]], username: str, output_path: str) -> Tuple[str, bool]:
        """
        Exports conversation to a highly styled, professional Word (.docx) file.
        Strictly adheres to corporate/academic standards (Times New Roman, margins, headers, footers).
        Handles PermissionError dynamically by writing to an alternative path.
        Returns: (actual_saved_path, is_new_file_created)
        """
        doc = Document()
        
        # 1. Page Margins Setup (A4 standard: Top/Bottom 2.5cm, Left 3.0cm, Right 2.0cm)
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        
        # 2. Configure Global Document Styles (Times New Roman 13pt, line spacing 1.3)
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Times New Roman'
        font_normal.size = Pt(13)
        
        # Line Spacing 1.3 lines
        style_normal.paragraph_format.line_spacing = 1.3
        style_normal.paragraph_format.space_after = Pt(6)
        
        # Heading 1: Section Title (Times New Roman 14pt, Bold, Upper Case)
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = 'Times New Roman'
        font_h1.size = Pt(14)
        font_h1.bold = True
        font_h1.color.rgb = docx.shared.RGBColor(0, 0, 0)
        style_h1.paragraph_format.space_before = Pt(12)
        style_h1.paragraph_format.space_after = Pt(12)
        style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading 2: Questions (Arial 13pt, Bold, Upper Case)
        style_h2 = doc.styles['Heading 2']
        font_h2 = style_h2.font
        font_h2.name = 'Arial'
        font_h2.size = Pt(13)
        font_h2.bold = True
        font_h2.color.rgb = docx.shared.RGBColor(94, 106, 210)  # Primary lavender color for accent
        style_h2.paragraph_format.space_before = Pt(16)
        style_h2.paragraph_format.space_after = Pt(6)
        
        # Heading 3: Answers Header (Arial 13pt, Bold, Regular Case)
        style_h3 = doc.styles['Heading 3']
        font_h3 = style_h3.font
        font_h3.name = 'Arial'
        font_h3.size = Pt(13)
        font_h3.bold = True
        font_h3.color.rgb = docx.shared.RGBColor(51, 65, 85)   # Charcoal
        style_h3.paragraph_format.space_before = Pt(10)
        style_h3.paragraph_format.space_after = Pt(4)
        
        # 3. Configure Header & Footer
        # Header (Right: Topic info, Times New Roman 11pt, Italic, border bottom)
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.paragraph_format.space_after = Pt(4)
        header_run = header_p.add_run("BÁO CÁO HỘI THOẠI WIKIBOT - HỆ THỐNG TRÍ TUỆ NHÂN TẠO")
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(11)
        header_run.font.italic = True
        header_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        add_p_border_bottom(header_p)
        
        # Footer (Left: Student Name & Exptime, Right: Page counter. Times New Roman 11pt, Italic)
        footer_p = section.footer.paragraphs[0]
        add_p_border_top(footer_p)
        footer_p.paragraph_format.space_before = Pt(4)
        # Add alignment tab at 16 cm (near the right margin)
        footer_p.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        
        export_time = datetime.now().strftime('%d/%m/%Y %H:%M')
        footer_run_left = footer_p.add_run(f"Người xuất: {username} | {export_time}\t")
        footer_run_left.font.name = 'Times New Roman'
        footer_run_left.font.size = Pt(11)
        footer_run_left.font.italic = True
        footer_run_left.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        
        footer_run_right = footer_p.add_run("Trang ")
        footer_run_right.font.name = 'Times New Roman'
        footer_run_right.font.size = Pt(11)
        footer_run_right.font.italic = True
        footer_run_right.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        add_page_number(footer_run_right)
        
        # 4. Write Document Content
        # Document Main Title
        title_p = doc.add_paragraph(f"BÁO CÁO HỘI THOẠI WIKIBOT", style='Heading 1')
        
        # Add metadata box
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_p.add_run(
            f"Tiêu đề cuộc trò chuyện: {conversation_title}\n"
            f"Ngày tạo: {export_time}\n"
            f"Số lượng tin nhắn: {len(messages)}"
        )
        meta_run.font.italic = True
        meta_run.font.size = Pt(10.5)
        meta_p.paragraph_format.space_after = Pt(24)
        
        # Write questions & answers
        for idx, msg in enumerate(messages):
            role = msg.get("role")
            content = msg.get("content", "").strip()
            
            if role == "user":
                # Question header
                q_p = doc.add_paragraph(f"CÂU HỎI {idx // 2 + 1}:", style='Heading 2')
                q_p.paragraph_format.keep_with_next = True
                
                # Question content
                content_p = doc.add_paragraph()
                # Indent user question a bit to distinguish
                content_p.paragraph_format.left_indent = Cm(0.5)
                run = content_p.add_run(content)
                run.font.bold = True
                content_p.paragraph_format.space_after = Pt(12)
            else:
                # Answer header
                a_p = doc.add_paragraph("TRẢ LỜI TỪ WIKIBOT:", style='Heading 3')
                a_p.paragraph_format.keep_with_next = True
                
                # Answer content - support markdown table conversion
                blocks = extract_content_blocks(content)
                for block in blocks:
                    if block["type"] == "text":
                        block_text = block["content"].strip()
                        if not block_text:
                            continue
                        # Standard paragraph
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(0.5)
                        
                        # Process quick inline formats (bold, italic, warning community)
                        lines = block_text.split('\n')
                        for l_idx, line in enumerate(lines):
                            if line.strip().startswith('⚠️ Cảnh báo:'):
                                # Display warning in red/orange bold
                                run = p.add_run(line + '\n' if l_idx < len(lines)-1 else line)
                                run.font.bold = True
                                run.font.color.rgb = docx.shared.RGBColor(245, 158, 11)  # Amber Warning
                            else:
                                p.add_run(line + '\n' if l_idx < len(lines)-1 else line)
                        p.paragraph_format.space_after = Pt(8)
                    else:
                        # Markdown Table Block
                        headers, rows = parse_markdown_table(block["content"])
                        if not headers:
                            continue
                            
                        # Add a table
                        num_cols = len(headers)
                        num_rows = len(rows) + 1
                        
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.style = 'Table Grid'
                        
                        # Style table headers
                        hdr_cells = table.rows[0].cells
                        for col_idx, text in enumerate(headers):
                            cell = hdr_cells[col_idx]
                            cell.text = text
                            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                            # Bold header
                            for run in cell.paragraphs[0].runs:
                                run.font.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                
                        # Style table rows
                        for r_idx, row_data in enumerate(rows):
                            row_cells = table.rows[r_idx + 1].cells
                            for col_idx, text in enumerate(row_data):
                                cell = row_cells[col_idx]
                                cell.text = text
                                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                                cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                                for run in cell.paragraphs[0].runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                                    
                        # Spacing after table
                        doc.add_paragraph().paragraph_format.space_after = Pt(8)
                
                # Seperator after complete dialogue pair
                if idx < len(messages) - 1:
                    sep_p = doc.add_paragraph()
                    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sep_run = sep_p.add_run("❖   ❖   ❖")
                    sep_run.font.color.rgb = docx.shared.RGBColor(200, 200, 200)
                    sep_p.paragraph_format.space_before = Pt(12)
                    sep_p.paragraph_format.space_after = Pt(12)
        
        # 5. Handle PermissionError & Save File
        is_new_file = False
        final_path = output_path
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            doc.save(output_path)
        except PermissionError:
            # File is locked (open in Word). Save to alternative name.
            base, ext = os.path.splitext(output_path)
            final_path = f"{base}_new{ext}"
            doc.save(final_path)
            is_new_file = True
            
        return final_path, is_new_file

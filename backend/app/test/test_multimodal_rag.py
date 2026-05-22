import os
import sys
import unittest
import shutil
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from unittest.mock import patch, MagicMock

# Thêm thư mục backend vào python path để import app
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from app.models.models import Base, User, Document, UserAISettings, Role
from app.services.document_processor import DocumentProcessor
from app.services.retriever import HybridRetriever

class TestMultimodalRAG(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # Thiết lập SQLite in-memory test database hoặc test.db cục bộ
        cls.engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(cls.engine)
        cls.Session = sessionmaker(bind=cls.engine)
        cls.db = cls.Session()
        
        # Tạo vai trò mặc định
        cls.public_role = Role(id=0, name="Public", level=99)
        cls.user_role = Role(id=3, name="Personal User", level=3)
        cls.db.add_all([cls.public_role, cls.user_role])
        cls.db.commit()
        
        # Tạo các người dùng thử nghiệm
        cls.user_a = User(id=1, username="user_a", full_name="User A", hashed_password="hash", role_id=3)
        cls.user_b = User(id=2, username="user_b", full_name="User B", hashed_password="hash", role_id=3)
        cls.db.add_all([cls.user_a, cls.user_b])
        cls.db.commit()
        
        # Tạo cấu hình AI cho người dùng
        cls.settings_a = UserAISettings(user_id=1, receive_community_knowledge=False)
        cls.settings_b_with_community = UserAISettings(user_id=2, receive_community_knowledge=True)
        cls.db.add_all([cls.settings_a, cls.settings_b_with_community])
        cls.db.commit()

        # Tạo một tài liệu PDF mẫu và Word mẫu để test trích xuất
        cls.backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        cls.pdf_path = os.path.join(cls.backend_dir, "data", "7425d760-eca9-413b-b518-d0f1b43ba268.pdf")
        cls.docx_path = os.path.join(cls.backend_dir, "data", "d23f22f6-33fe-4a50-ad4e-72b373aa3aec.docx")
        
        # Tạo file Word mẫu có chứa bảng biểu thật sự bằng python-docx để test trích xuất Word
        try:
            import docx
            doc = docx.Document()
            doc.add_paragraph("Quy định Nghỉ phép Công ty WikiBot")
            # Thêm bảng biểu thực tế
            table = doc.add_table(rows=3, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cấp bậc'
            hdr_cells[1].text = 'Số ngày phép'
            hdr_cells[2].text = 'Ghi chú'
            row_cells = table.rows[1].cells
            row_cells[0].text = 'Nhân viên'
            row_cells[1].text = '15'
            row_cells[2].text = 'Nghỉ tiêu chuẩn'
            row_cells2 = table.rows[2].cells
            row_cells2[0].text = 'Quản lý'
            row_cells2[1].text = '18'
            row_cells2[2].text = 'Nghỉ nâng cao'
            os.makedirs(os.path.dirname(cls.docx_path), exist_ok=True)
            doc.save(cls.docx_path)
        except Exception as e:
            print(f"Không thể tạo file Word mẫu cho test: {e}")
        
        # Khởi tạo DocumentProcessor dùng db in-memory
        # (Lưu ý: DocumentProcessor thực tế có thể gọi get_embedding_model, 
        #  trong unittest này chúng ta muốn đảm bảo các hàm trích xuất đa phương thức hoạt động đúng)
        cls.processor = DocumentProcessor(db_session=cls.db)
        
        # Sử dụng collection riêng cho Unit Test để tránh xung đột chiều (dimension 384 vs 2048) với database thật
        try:
            cls.processor.chroma_client.delete_collection("documents_test_rag")
        except Exception:
            pass
        cls.processor.collection = cls.processor.chroma_client.get_or_create_collection(
            name="documents_test_rag",
            metadata={"hnsw:space": "cosine"}
        )

    @classmethod
    def tearDownClass(cls):
        cls.db.close()
        # Xóa collection test để giải phóng tài nguyên và dọn dẹp db
        try:
            cls.processor.chroma_client.delete_collection("documents_test_rag")
        except Exception:
            pass
        # Dọn dẹp thư mục extracted_images nếu có sinh ra
        images_dir = os.path.join(cls.backend_dir, "data", "extracted_images")
        if os.path.exists(images_dir):
            # Không xóa hoàn toàn nếu đang có dữ liệu thật, chỉ xóa các file test sinh ra
            for f in os.listdir(images_dir):
                if f.startswith("doc_9999_"):
                    try:
                        os.remove(os.path.join(images_dir, f))
                    except Exception:
                        pass

    @patch('pdfplumber.open')
    @patch('app.services.vision_processor.VisionProcessor.extract_text_from_image')
    def test_1_extract_pdf_multimodal(self, mock_extract_text, mock_pdfplumber_open):
        """Kiểm tra trích xuất tài liệu PDF chứa bảng biểu và hình ảnh"""
        print("\n--- TEST: Trich xuat PDF da phuong thuc ---")
        if not os.path.exists(self.pdf_path):
            self.skipTest(f"Không tìm thấy file PDF mẫu tại {self.pdf_path}")
            
        # 1. Mock pdfplumber để giả lập trích xuất bảng biểu
        mock_pdf = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Phiếu mua hàng"
        
        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["Tên sản phẩm", "Số lượng", "Đơn giá"],
            ["Bàn phím cơ", "1", "1.500.000"],
            ["Chuột không dây", "2", "800.000"]
        ]
        mock_page.find_tables.return_value = [mock_table]
        mock_pdf.pages = [mock_page]
        mock_pdfplumber_open.return_value.__enter__.return_value = mock_pdf
        
        # 2. Mock VisionProcessor OCR
        mock_extract_text.return_value = "HOÁ ĐƠN BÁN HÀNG"
        
        # 3. Mock fitz (PyMuPDF) để giả lập trích xuất ảnh nhúng PNG hợp lệ (sinh bằng Pillow)
        from PIL import Image
        import io
        img = Image.new('RGB', (10, 10), color='red')
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        valid_png_bytes = img_byte_arr.getvalue()
        
        with patch('fitz.open') as mock_fitz_open:
            mock_fitz_doc = MagicMock()
            mock_fitz_page = MagicMock()
            
            # get_images trả về danh sách ảnh: [(xref, smask, width, height, bpc, colorspace, ...)]
            mock_fitz_page.get_images.return_value = [(123, 0, 10, 10, 8, 'DeviceRGB', '', '', 0)]
            mock_fitz_doc.__getitem__.return_value = mock_fitz_page
            mock_fitz_doc.close.return_value = None
            
            # extract_image trả về dict chứa bytes ảnh nhúng hợp lệ
            mock_fitz_doc.extract_image.return_value = {
                "image": valid_png_bytes,
                "ext": "png"
            }
            mock_fitz_open.return_value = mock_fitz_doc
            
            # Gọi hàm trích xuất đa phương thức nâng cao
            elements = self.processor._extract_pdf_multimodal(self.pdf_path, document_id=9999)
            
        self.assertIsNotNone(elements)
        self.assertGreater(len(elements), 0)
        
        # Kiểm tra xem có chứa cấu trúc bảng biểu Markdown được bóc tách hay không
        has_table = any(el.get("type") == "table" or "|" in el.get("content", "") for el in elements)
        self.assertTrue(has_table, "PDF trích xuất phải phát hiện và chuyển đổi được bảng Markdown")
        
        # Kiểm tra xem có hình ảnh nào được trích xuất ra thư mục hay không
        images_dir = os.path.join(self.backend_dir, "data", "extracted_images")
        has_extracted_images = False
        if os.path.exists(images_dir):
            has_extracted_images = any(f.startswith("doc_9999_") for f in os.listdir(images_dir))
        
        self.assertTrue(has_extracted_images, "Phải trích xuất được hình ảnh nhúng từ PDF ra thư mục dữ liệu")
        print("=> PDF trich xuat thanh cong cac phan doan va bang bieu Markdown!")

    def test_2_extract_docx_multimodal(self):
        """Kiểm tra trích xuất tài liệu Word (.docx) chứa bảng biểu và hình ảnh"""
        print("\n--- TEST: Trich xuat Word (.docx) da phuong thuc ---")
        if not os.path.exists(self.docx_path):
            self.skipTest(f"Không tìm thấy file Word mẫu tại {self.docx_path}")
            
        # Gọi hàm trích xuất đa phương thức nâng cao cho file Word
        elements = self.processor._extract_docx_multimodal(self.docx_path, document_id=9999)
        
        self.assertIsNotNone(elements)
        self.assertGreater(len(elements), 0)
        
        # Xác nhận trích xuất được bảng biểu Word thành Markdown Table
        has_table = any(el.get("type") == "table" or "|" in el.get("content", "") for el in elements)
        self.assertTrue(has_table, "Word trích xuất phải chuyển đổi được bảng Word thành Markdown Table")
        print("=> Word trich xuat thanh cong cac doan van va bang Markdown!")

    def test_3_global_cross_tenant_retrieval(self):
        """Kiểm tra truy xuất vector chéo cộng đồng (Global Cross-Tenant RAG)"""
        print("\n--- TEST: Truy xuat Vector cheo cong dong ---")
        
        # 1. Tạo tài liệu của User A được đánh dấu chia sẻ cộng đồng
        doc_a = Document(
            id=101,
            filename="user_a_doc.pdf",
            original_name="Tài liệu chia sẻ cộng đồng.pdf",
            file_path=self.pdf_path,
            file_type="pdf",
            uploaded_by=self.user_a.id,
            role_id=3,
            is_public_community=True, # Đánh dấu chia sẻ cộng đồng
            privacy_mode=False
        )
        self.db.add(doc_a)
        self.db.commit()
        
        # Giả lập nạp các chunks vào vector DB ChromaDB
        # Để đảm bảo test này độc lập, chúng ta sẽ add trực tiếp vài chunks giả lập vào ChromaDB collection
        collection = self.processor.collection
        
        chunk_id_1 = "doc_101_chunk_0"
        chunk_content_1 = "Tri thức cộng đồng từ User A: Công thức RAG nâng cấp WikiBot năm 2026."
        chunk_embedding_1 = self.processor.embedding_model.encode([chunk_content_1])
        if not isinstance(chunk_embedding_1, list):
            chunk_embedding_1 = chunk_embedding_1.tolist()
            
        # Xóa chunk cũ nếu có để tránh trùng lặp
        try:
            collection.delete(ids=[chunk_id_1])
        except Exception:
            pass
            
        collection.add(
            documents=[chunk_content_1],
            metadatas=[{
                "source": "Tài liệu chia sẻ cộng đồng.pdf",
                "document_id": 101,
                "chunk_index": 0,
                "role_id": 3,
                "file_type": "pdf",
                "element_type": "narrative",
                "is_public_community": True, # Metadata cộng đồng chéo
                "uploaded_by": self.user_a.id
            }],
            ids=[chunk_id_1],
            embeddings=chunk_embedding_1
        )
        
        # 2. Truy vấn dưới tư cách User B (có bật receive_community_knowledge = True)
        # User B có accessible_role_ids = [3, 0] (vai trò Personal User và Public)
        accessible_roles = [3, 0]
        
        # Khởi tạo HybridRetriever
        retriever = HybridRetriever(self.processor)
        
        # Thực hiện tìm kiếm tương tự
        # Chúng ta sẽ nâng cấp hàm search_similar để lọc filter $or nếu có bật cộng đồng
        # Giả lập hành vi thông qua việc gọi search với cờ receive_community_knowledge = True
        # Ở đây ta sẽ gọi search_similar của processor để kiểm tra logic lọc ở tầng vector DB
        results = self.processor.search_similar(
            query="WikiBot năm 2026",
            accessible_role_ids=accessible_roles,
            top_k=5,
            max_distance=1.0, # Cho khoảng cách lớn để dễ khớp
            receive_community=True # Cờ nhận tri thức cộng đồng
        )
        
        # Xác nhận tìm được tài liệu cộng đồng của User A
        found_community_doc = any(r["metadata"].get("document_id") == 101 for r in results)
        self.assertTrue(found_community_doc, "User B phải truy xuất được tài liệu cộng đồng của User A khi bật receive_community")
        
        # 3. Truy vấn dưới tư cách User B nhưng TẮT nhận cộng đồng
        results_no_community = self.processor.search_similar(
            query="WikiBot năm 2026",
            accessible_role_ids=accessible_roles,
            top_k=5,
            max_distance=1.0,
            receive_community=False
        )
        
        found_community_doc_no = any(r["metadata"].get("document_id") == 101 for r in results_no_community)
        self.assertFalse(found_community_doc_no, "User B tuyệt đối không được truy xuất tài liệu cộng đồng của User A khi tắt receive_community")
        
        # Dọn dẹp ChromaDB sau khi test xong
        try:
            collection.delete(ids=[chunk_id_1])
        except Exception:
            pass
            
        print("=> Logic truy xuat vector cheo cong dong duoc xac thuc bao mat va chinh xac!")

if __name__ == "__main__":
    unittest.main()
